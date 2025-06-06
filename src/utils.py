"""
Utility functions for the RAG system frontend.
"""
from datetime import datetime
from typing import List, Tuple, Union, Optional
from docx import Document as DocxDocument
from docx.shared import Pt
import os


def render_to_docx(
  data: Union[Tuple[str, str], List[Tuple[str, str]]], 
  output_path: str = "adidas_report.docx"
) -> str:
  """Create or append to a DOCX report with query and answer data.
  
  Args:
    data: Either a single (query, answer) tuple or list of (query, answer) tuples
    output_path: Path to save the document
    
  Returns:
    Path to the saved document
  """
  # Check if file exists and load it, otherwise create new document
  if os.path.exists(output_path):
    doc = DocxDocument(output_path)
    # Add separator before new content
    doc.add_paragraph("═" * 60)
    doc.add_paragraph()
  else:
    doc = DocxDocument()
    # Title for new document
    title = doc.add_heading("Adidas Annual Report Analysis", level=1)
    title.alignment = 1  # Center align

  # Add timestamp
  doc.add_paragraph(f"Session: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", style='Intense Quote')

  # Handle single tuple or list of tuples
  if isinstance(data, tuple):
    # Single Q&A pair
    query_text, answer_text = data
    _add_qa_section(doc, query_text, answer_text)
  else:
    # Multiple Q&A pairs
    for i, (query_text, answer_text) in enumerate(data, 1):
      _add_qa_section(doc, query_text, answer_text, question_num=i)
      # Add separator except for last item
      if i < len(data):
        doc.add_paragraph("─" * 50)

  # Save document
  doc.save(output_path)
  return output_path


def _add_qa_section(doc, query_text: str, answer_text: str, question_num: Optional[int] = None) -> None:
  """Add a question and answer section to the document."""
  # Question section
  if question_num:
    doc.add_heading(f"Question {question_num}:", level=2)
  else:
    doc.add_heading("Question:", level=2)
    
  q_para = doc.add_paragraph(query_text)
  if q_para.style and hasattr(q_para.style, 'font'):
    q_para.style.font.size = Pt(12)

  # Answer section  
  if question_num:
    doc.add_heading(f"Answer {question_num}:", level=2)
  else:
    doc.add_heading("Answer:", level=2)
  
  # Body content
  paragraphs = answer_text.split('\n\n')
  for para in paragraphs:
    if para.strip():
      p = doc.add_paragraph(para.strip())
      if p.style and hasattr(p.style, 'font'):
        p.style.font.size = Pt(11)


def query_rag_system_streamlit(rag_chain, query_text: str) -> str:
  """Process a query through the RAG system for Streamlit interface."""
  try:
    answer = rag_chain.invoke(query_text)
    return answer
  except Exception as e:
    error_msg = f"Error processing query: {str(e)}"
    return error_msg 